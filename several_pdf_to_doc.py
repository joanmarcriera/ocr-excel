#!/usr/bin/env python3

import pytesseract
import cv2
import pandas as pd
import PIL
from openpyxl import Workbook
import os
import shutil
import datetime
from docx import Document


# Import other OCR libraries
try:
    import pyocr
    from pyocr import builders
    pyocr_available = True
except ImportError:
    pyocr_available = False

try:
    import textract
    textract_available = True
except ImportError:
    textract_available = False

def ensure_path(path):
    if not os.path.exists(path):
      os.makedirs(path)

working_path='./temp'
output_path='./output'

ensure_path(working_path)
ensure_path(output_path)

from tkinter import filedialog
from tkinter import *

root = Tk()
root.filename =  filedialog.askopenfilename(initialdir = "/", title = "Select file", filetypes = (("jpeg files", "*.jpg"),("PDF files", "*.pdf"), ("all files", "*.*")))
print (root.filename)

# Get current timestamp
now = datetime.datetime.now()
timestamp = now.strftime("%Y%m%d%H%M%S")

# Get file name and extension
file_name, file_ext = os.path.splitext(os.path.basename(root.filename))

# Create new file name with timestamp
new_file_name = f'{file_name}_{timestamp}{file_ext}'

# Construct the destination path
destination_path = os.path.join(working_path, new_file_name)

# Copy the file to the destination path
shutil.copy2(root.filename, destination_path)

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO

def pdf_to_text(file_path):
    with open(file_path, 'rb') as file:
        # Create a PDF resource manager
        rsrcmgr = PDFResourceManager()

        # Create a StringIO object
        retstr = StringIO()

        # Create a PDF device object
        device = TextConverter(rsrcmgr, retstr, laparams=LAParams())

        # Create a PDF interpreter object
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # Process each page contained in the PDF document
        for page in PDFPage.get_pages(file):
            interpreter.process_page(page)

        # Get the text from the StringIO object
        text = retstr.getvalue()

        # Close the StringIO object and the PDF device
        device.close()
        retstr.close()

        return text

text = pdf_to_text(destination_path)
print(text)


from docx import Document

# Create a new Word document
document = Document()

import re
import string

# You can also use str.translate method to remove non-ascii characters from the text
text = text.translate(str.maketrans('', '', string.punctuation))

# Use regular expression to remove non-alphanumeric characters
text = re.sub(r'[^\w\s]', '', text)

text = ''.join(c for c in text if c.isprintable())

text = text.encode('ascii', 'ignore').decode()

# Add the filtered text to the document
document.add_paragraph(text)

# Add the extracted text to the document
document.add_paragraph(text)


# Create new file name with timestamp
output_file_name = f'{file_name}_{timestamp}.docx'

# Construct the destination path
output_destination_path = os.path.join(output_path, output_file_name)

# Save the document
document.save(output_destination_path)
